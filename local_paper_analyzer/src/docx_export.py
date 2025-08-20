import re
from typing import List, Tuple, Optional
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def _add_code_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    font = run.font
    font.name = 'Consolas'
    font.size = Pt(10)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '等线')

def _add_formatted_paragraph(doc: Document, text: str, style: Optional[str] = None):
    p = doc.add_paragraph()
    if style:
        p.style = style
    pos = 0
    pattern = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = p.add_run(token[2:-2]); run.bold = True
        elif token.startswith("*"):
            run = p.add_run(token[1:-1]); run.italic = True
        elif token.startswith("`"):
            run = p.add_run(token[1:-1]); run.font.name = "Consolas"; run.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])

def markdown_to_docx(md_text: str, out_path: str):
    doc = Document()
    in_code = False
    code_buf: List[str] = []
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        fence = re.match(r"^```(\w+)?\s*$", line)
        if fence:
            if in_code:
                _add_code_paragraph(doc, "\n".join(code_buf)); code_buf = []; in_code = False
            else:
                in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        if '|' in line and line.strip().startswith('|'):
            tbl_lines = []
            j = i
            while j < len(lines) and '|' in lines[j]:
                tbl_lines.append(lines[j]); j += 1
            rows = []
            for tl in tbl_lines:
                rows.append([c.strip() for c in tl.strip().strip('|').split('|')])
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                for r_idx, r_vals in enumerate(rows):
                    for c_idx in range(ncols):
                        txt = r_vals[c_idx] if c_idx < len(r_vals) else ""
                        cell = table.cell(r_idx, c_idx)
                        run = cell.paragraphs[0].add_run(txt)
                        if r_idx == 0:
                            run.bold = True
            i = j; continue

        if line.startswith("###### "): doc.add_heading(line[7:], level=4); i += 1; continue
        if line.startswith("##### "): doc.add_heading(line[6:], level=4); i += 1; continue
        if line.startswith("#### "):  doc.add_heading(line[5:], level=3); i += 1; continue
        if line.startswith("### "):   doc.add_heading(line[4:], level=2); i += 1; continue
        if line.startswith("## "):    doc.add_heading(line[3:], level=1); i += 1; continue
        if line.startswith("# "):     doc.add_heading(line[2:], level=0); i += 1; continue

        if re.match(r"^\s*-\s+", line):
            _add_formatted_paragraph(doc, re.sub(r"^\s*-\s+", "", line), style="List Bullet"); i += 1; continue
        if re.match(r"^\s*\d+\.\s+", line):
            _add_formatted_paragraph(doc, re.sub(r"^\s*\d+\.\s+", "", line), style="List Number"); i += 1; continue

        if line.strip() == "---":
            doc.add_paragraph("-" * 20); i += 1; continue
        if line.strip() == "":
            doc.add_paragraph(""); i += 1; continue

        line = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", line)
        _add_formatted_paragraph(doc, line); i += 1

    if in_code and code_buf:
        _add_code_paragraph(doc, "\n".join(code_buf))

    doc.save(out_path)
